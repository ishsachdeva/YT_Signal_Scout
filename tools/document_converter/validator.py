"""Warning-only validation of generated Markdown."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document


@dataclass(frozen=True)
class ValidationResult:
    """Validation warnings associated with one output file."""

    warnings: tuple[str, ...]


def validate_conversion(source_path: Path, markdown_path: Path) -> ValidationResult:
    """Validate output presence, content, headings, and source tables."""
    warnings: list[str] = []

    if not markdown_path.is_file():
        return ValidationResult((f"Markdown file was not created: {markdown_path.name}",))

    markdown = markdown_path.read_text(encoding="utf-8")
    if not markdown.strip():
        warnings.append(f"Markdown file is empty: {markdown_path.name}")

    if not re.search(r"(?m)^#{1,6}\s+\S", markdown):
        warnings.append(f"No Markdown headings detected: {markdown_path.name}")

    source = Document(source_path)
    markdown_table_count = _count_markdown_tables(markdown)
    if source.tables and markdown_table_count < len(source.tables):
        warnings.append(
            f"Source contains {len(source.tables)} table(s), but only "
            f"{markdown_table_count} Markdown table(s) were detected: "
            f"{markdown_path.name}"
        )

    return ValidationResult(tuple(warnings))


def _count_markdown_tables(markdown: str) -> int:
    separator_cell = re.compile(r"\s*:?-{3,}:?\s*")
    count = 0
    for line in markdown.splitlines():
        if "|" not in line:
            continue
        cells = line.strip().strip("|").split("|")
        if cells and all(separator_cell.fullmatch(cell) for cell in cells):
            count += 1
    return count
